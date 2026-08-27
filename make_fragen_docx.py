#!/usr/bin/env python3
"""Erzeugt ein DOCX mit offenen Fragen an den Kunden (zum Optimieren/Weitergeben)."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ANTHRAZIT = RGBColor(0x2D, 0x2C, 0x2E)
GOLD = RGBColor(0x8C, 0x73, 0x45)

doc = Document()

# Grundschrift
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = ANTHRAZIT
    return p

def h2(text):
    p = doc.add_paragraph(); p.space_before = Pt(12)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = GOLD
    return p

def q(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)

def answer():
    p = doc.add_paragraph()
    r = p.add_run("Antwort: "); r.italic = True; r.font.color.rgb = RGBColor(0x6B, 0x6A, 0x6D)
    p.add_run("________________________________________________")

def note(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x6B, 0x6A, 0x6D)

# ---------- Titel ----------
h1("Bauwelt Handwerk – Website")
p = doc.add_paragraph()
r = p.add_run("Offene Fragen & benötigte Informationen"); r.bold = True; r.font.size = Pt(13)
note("Bitte ausfüllen bzw. mit dem Kunden klären. Nichts davon blockiert die Vorschau – "
     "es geht darum, die Seite Schritt für Schritt „scharf“ zu schalten. "
     "Vorschau: https://cansi798.github.io/Bauwelt/")

# ---------- E-Mail ----------
h2("1. E-Mail-Adressen")
note("Empfehlung: eine zentrale Sammeladresse einrichten, dahinter kann intern verteilt werden.")
q("info@bauwelt-handwerk.de – zentrale Adresse für Anfragen über die Website (EMPFOHLEN). "
  "Soll diese eingerichtet werden? Ja/Nein")
answer()
q("Welche Adresse soll auf der Website als Kontakt stehen? "
  "(aktuell: j.iqbal@bauwelt-handwerk.de – oder info@ / anfrage@ …)")
answer()
q("Separate Adresse für Bewerbungen gewünscht? (z. B. bewerbung@bauwelt-handwerk.de) – "
  "sonst gehen Bewerbungen an die zentrale Adresse")
answer()
q("Separate Adresse für Kooperations-/Partneranfragen (B2B)? (z. B. partner@bauwelt-handwerk.de)")
answer()
q("Adresse für Datenschutz-Anfragen (DSGVO)? (z. B. datenschutz@bauwelt-handwerk.de) – optional")
answer()

# ---------- Anfragen/Formular ----------
h2("2. Anfragen & Formular")
note("Aktuell öffnen die Formulare das E-Mail-Programm mit vorausgefüllter Anfrage (Übergangslösung).")
q("Soll es dabei bleiben, oder ein echtes Formular-Backend, das Anfragen ohne Mailprogramm "
  "versendet? (z. B. Formspree/Web3Forms – kostenlos für kleine Mengen)")
answer()
q("An welche Adresse sollen Formular-Anfragen konkret gehen?")
answer()

# ---------- Terminbuchung ----------
h2("3. Online-Terminbuchung")
q("Welcher Anbieter soll eingebettet werden? (z. B. Calendly, Cal.com) – oder vorerst nur "
  "Telefon/WhatsApp?")
answer()
q("Falls Anbieter: bitte Zugang/Link nennen, dann wird der Kalender eingebettet.")
answer()

# ---------- Preise ----------
h2("4. Badkalkulator – echte Preise")
note("Aktuell Schätz-Platzhalter. Bitte reale Richtwerte, damit der Rechner stimmt.")
q("Grundpauschale (fixe Basis, z. B. Demontage/Planung): __________ €")
q("Preis pro m² – Standard: ______ €   Komfort: ______ €   Premium: ______ €")
q("Aufpreise – Bodengleiche Dusche: ______ €   Fußbodenheizung: ______ €   Barrierefrei: ______ €")
q("Gewünschte Preisspanne (z. B. ± 10 %): ______ %")

# ---------- Inhalte ----------
h2("5. Inhalte, die echt werden müssen")
q("Zahlen & Fakten (aktuell Beispiele): Anzahl sanierte Bäder, Jahre Erfahrung, Gewährleistung … "
  "Bitte nur belegbare Zahlen.")
answer()
q("Kundenstimmen: 3 echte Zitate mit Name/Kürzel, Ort, Sternebewertung – und Einverständnis "
  "zur Veröffentlichung.")
answer()
q("Stimmen die 6 Leistungsbereiche? (Bad, Modernisierung, Dach & Fassade, Innenausbau, "
  "Sanitär/Heizung/Elektro, Maler & Böden) – fehlt/entfällt etwas?")
answer()
q("Einzugsgebiet: „Hamburg und Umgebung / rund um Norderstedt“ korrekt? Oder konkrete Orte/Umkreis?")
answer()
q("Erreichbarkeit „Mo–Fr, 8–18 Uhr“ korrekt?")
answer()

# ---------- Bilder ----------
h2("6. Bilder")
q("Porträtfoto von Jamil Iqbal (Ansprechpartner) – für den Kontaktbereich. Vorhanden?")
answer()
q("Team-/Über-uns-Foto (echt, in Arbeitskleidung) – vorhanden?")
answer()
note("Weitere Projektfotos jederzeit möglich. Namensvorgaben und Formate siehe Datei BILDER.md.")

# ---------- Rechtliches ----------
h2("7. Rechtliches")
q("Handwerkskammer-Pflichtangaben fürs Impressum: zuständige Handwerkskammer, Berufsbezeichnung, "
  "ggf. Berufshaftpflicht – welche Angaben müssen rein?")
answer()
q("Datenschutzerklärung: ist ein Entwurf – wird sie vor Livegang geprüft (GF/Datenschutzberater)?")
answer()

# ---------- Kanäle ----------
h2("8. Kanäle & Domain")
q("WhatsApp-Nummer für den Chat-Button korrekt? (aktuell 0151 56156578)")
answer()
q("Social-Media-Profile (Instagram/Facebook) zum Verlinken vorhanden?")
answer()
q("Eigene Domain (www.bauwelt-handwerk.de) anbinden? Falls ja: Zugang zu den DNS-Einstellungen?")
answer()

# ---------- Sonstiges ----------
h2("9. Sonstiges")
q("Anrede: durchgängig „Sie“ (aktuell, CI-konform) – oder „du“ gewünscht?")
answer()
q("Weitere Wünsche / Anmerkungen zum Aufbau oder Design?")
answer()

doc.save("Offene-Fragen-Kunde-Bauwelt.docx")
print("Offene-Fragen-Kunde-Bauwelt.docx erstellt")
